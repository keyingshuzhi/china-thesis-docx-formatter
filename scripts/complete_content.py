#!/usr/bin/env python3
"""Find thesis placeholders and apply only user-reviewed replacements."""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document

PLACEHOLDER = re.compile(r"\[\[(?:TODO|待补充)[:：][^\]]+\]\]|\[(?:TODO|待补充)[:：][^\]]+\]|【(?:待补充|待填写)[^】]*】|_{4,}")


def paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def find_placeholders(document: Document) -> list[dict]:
    found = []
    for index, paragraph in enumerate(paragraphs(document), start=1):
        for match in PLACEHOLDER.finditer(paragraph.text):
            found.append({"paragraph": index, "placeholder": match.group(), "context": paragraph.text[:240]})
    return found


def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> tuple[int, list[str]]:
    applied, skipped = 0, []
    for placeholder, replacement in replacements.items():
        matching_runs = [run for run in paragraph.runs if placeholder in run.text]
        if matching_runs:
            for run in matching_runs:
                run.text = run.text.replace(placeholder, replacement)
                applied += 1
        elif placeholder in paragraph.text:
            skipped.append(placeholder)
    return applied, skipped


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("--report", type=Path, help="Write placeholders and context as JSON.")
    parser.add_argument("--replacements", type=Path, help="Reviewed JSON object: {placeholder: replacement}.")
    parser.add_argument("--output", type=Path, help="Required when applying replacements.")
    args = parser.parse_args()
    document = Document(args.input_docx)
    found = find_placeholders(document)
    if args.report:
        args.report.write_text(json.dumps({"source": args.input_docx.name, "placeholders": found}, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        print(f"Wrote {len(found)} placeholder(s) to {args.report}")
    else:
        print(json.dumps(found, ensure_ascii=False, indent=2))
    if args.replacements:
        if not args.output:
            parser.error("--output is required with --replacements")
        replacements = json.loads(args.replacements.read_text(encoding="utf-8"))
        if not isinstance(replacements, dict) or not all(isinstance(key, str) and isinstance(value, str) for key, value in replacements.items()):
            parser.error("Replacement JSON must be an object mapping placeholder strings to text.")
        applied, skipped = 0, []
        for paragraph in paragraphs(document):
            count, unresolved = replace_in_paragraph(paragraph, replacements)
            applied += count
            skipped.extend(unresolved)
        document.save(args.output)
        print(f"Applied {applied} replacement(s): {args.output}")
        if skipped:
            print("Skipped placeholders split across runs: " + ", ".join(sorted(set(skipped))))


if __name__ == "__main__":
    main()
