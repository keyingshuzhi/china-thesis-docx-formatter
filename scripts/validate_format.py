#!/usr/bin/env python3
"""Validate template-derived rules and optionally audit a formatted DOCX copy."""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document

from docx_rules import DEGREE_PROFILES, RULES_SCHEMA_VERSION, SECTION_LABELS, _east_asia_font, _optional_cm, paragraph_role, read_rules, role_format, semantic_section_role


def validate_rules(rules: dict) -> list[str]:
    errors: list[str] = []
    required = {"schema_version", "source", "page", "styles", "role_rules", "table_rules", "extraction_warnings"}
    missing = required - rules.keys()
    if missing:
        errors.append(f"Missing rule keys: {', '.join(sorted(missing))}")
    if rules.get("schema_version") != RULES_SCHEMA_VERSION:
        errors.append(f"Unsupported schema_version: {rules.get('schema_version')!r}")
    if not isinstance(rules.get("source"), dict) or not rules.get("source", {}).get("template_filename"):
        errors.append("source.template_filename is required")
    if not isinstance(rules.get("page"), dict):
        errors.append("page must be an object")
    if not isinstance(rules.get("role_rules"), dict) or "body" not in rules.get("role_rules", {}):
        errors.append("role_rules.body is required")
    profile = rules.get("degree_profile")
    if profile is not None and (not isinstance(profile, dict) or profile.get("level") not in DEGREE_PROFILES):
        errors.append("degree_profile.level must be auto, undergraduate, master, or doctoral")
    for role, rule in rules.get("role_rules", {}).items():
        if not isinstance(rule, dict) or not isinstance(rule.get("format"), dict):
            errors.append(f"role_rules.{role}.format must be an object")
    return errors


def _matches(expected: float | None, actual: float | None, tolerance: float = 0.02) -> bool:
    return expected is None or (actual is not None and abs(expected - actual) <= tolerance)


def audit_document(document: Document, rules: dict) -> list[str]:
    findings: list[str] = []
    page = rules.get("page", {})
    if document.sections:
        section = document.sections[0]
        for key, actual in (("margin_top_cm", _optional_cm(section.top_margin)), ("margin_bottom_cm", _optional_cm(section.bottom_margin)),
                            ("margin_left_cm", _optional_cm(section.left_margin)), ("margin_right_cm", _optional_cm(section.right_margin))):
            if not _matches(page.get(key), actual):
                findings.append(f"Page {key} is {actual} cm; expected {page[key]} cm.")
    for number, paragraph in enumerate(document.paragraphs, start=1):
        role = paragraph_role(paragraph, rules)
        expected = role_format(rules, role) if role else None
        if not expected or not paragraph.text.strip():
            continue
        font = expected.get("font", {})
        for run in paragraph.runs:
            if font.get("size_pt") is not None and (run.font.size is None or abs(run.font.size.pt - font["size_pt"]) > 0.05):
                findings.append(f"Paragraph {number} ({role}) has a font-size mismatch.")
                break
            expected_face = font.get("east_asia") or font.get("name")
            actual_face = _east_asia_font(run._element) or run.font.name
            if expected_face and actual_face != expected_face:
                findings.append(f"Paragraph {number} ({role}) has a font-family mismatch.")
                break
    return findings


def audit_degree_structure(document: Document, degree: str) -> list[str]:
    found = {semantic_section_role(paragraph.text) for paragraph in document.paragraphs}
    return [f"{SECTION_LABELS[role]} was not detected; confirm whether this school requires it."
            for role in DEGREE_PROFILES[degree]["review_sections"] if role not in found]

def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("rules", type=Path)
    parser.add_argument("--document", type=Path, help="Formatted thesis copy to audit against the rules.")
    parser.add_argument("--degree", choices=("undergraduate", "master", "doctoral"),
                        help="Report common Chinese degree-thesis sections that need manual confirmation.")
    parser.add_argument("--strict-degree", action="store_true",
                        help="Treat undetected degree-profile sections as validation failures; use only after confirming the school requires them.")
    args = parser.parse_args()
    rules = read_rules(args.rules)
    errors = validate_rules(rules)
    if errors:
        print("Rules are invalid:", *[f"- {error}" for error in errors], sep="\n")
        raise SystemExit(1)
    print(f"Rules for '{rules['source']['template_filename']}' are valid.")
    if args.document:
        document = Document(args.document)
        findings = audit_document(document, rules)
        degree_findings = audit_degree_structure(document, args.degree) if args.degree else []
        if degree_findings:
            print("Degree-structure review:", *[f"- {item}" for item in degree_findings], sep="\n")
        if args.strict_degree:
            findings.extend(degree_findings)
        if findings:
            print("Document does not match all audited rules:", *[f"- {item}" for item in findings], sep="\n")
            raise SystemExit(1)
        print(f"Document '{args.document.name}' matches the audited rules.")


if __name__ == "__main__":
    main()
