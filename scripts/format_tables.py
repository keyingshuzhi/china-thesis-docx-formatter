#!/usr/bin/env python3
"""Apply observed table formatting from an extracted thesis-template JSON."""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

from docx_rules import apply_format, read_rules


def format_tables(document: Document, rules: dict) -> int:
    rule = rules.get("table_rules", {})
    if not rule.get("observed"):
        return 0
    for table in document.tables:
        alignment = rule.get("alignment")
        if alignment and hasattr(WD_TABLE_ALIGNMENT, alignment):
            table.alignment = getattr(WD_TABLE_ALIGNMENT, alignment)
        if rule.get("style_name"):
            try:
                table.style = rule["style_name"]
            except (KeyError, ValueError):
                pass
        for row in table.rows:
            for cell in row.cells:
                vertical = rule.get("cell_vertical_alignment")
                if vertical and hasattr(WD_CELL_VERTICAL_ALIGNMENT, vertical):
                    cell.vertical_alignment = getattr(WD_CELL_VERTICAL_ALIGNMENT, vertical)
                for paragraph in cell.paragraphs:
                    apply_format(paragraph, rule.get("cell_text_format"))
    return len(document.tables)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--rules", required=True, type=Path)
    args = parser.parse_args()
    document = Document(args.input_docx)
    count = format_tables(document, read_rules(args.rules))
    document.save(args.output_docx)
    print(f"Formatted {count} table(s): {args.output_docx}")


if __name__ == "__main__":
    main()
