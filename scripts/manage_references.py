#!/usr/bin/env python3
"""Audit numeric citations and repair numbering from user-supplied reference data."""
from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from pathlib import Path

from docx import Document

from docx_rules import paragraph_role, read_rules

# Both half-width and full-width brackets are encountered in domestic thesis
# drafts.  Keep the original bracket form when only renumbering labels.
BRACKETED = re.compile(r"([\[［])([0-9,，、\-–\s]+)([\]］])")
ENTRY = re.compile(r"^\s*[\[［](\d+)[\]］]\s*(.+)$")


def all_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def cited_numbers(text: str) -> set[int]:
    result: set[int] = set()
    for _, group, _ in BRACKETED.findall(text):
        for token in re.findall(r"\d+", group):
            result.add(int(token))
    return result


def reference_entries(document: Document, rules: dict) -> list[dict]:
    active, entries = False, []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        if paragraph_role(paragraph, rules) == "reference_heading":
            active = True
            continue
        if not active:
            continue
        match = ENTRY.match(paragraph.text)
        if match:
            entries.append({"paragraph": index, "number": int(match.group(1)), "text": match.group(2)})
    return entries


def audit(document: Document, rules: dict) -> dict:
    citations = sorted({number for paragraph in all_paragraphs(document) for number in cited_numbers(paragraph.text)})
    entries = reference_entries(document, rules)
    numbers = [entry["number"] for entry in entries]
    counts = Counter(numbers)
    return {
        "citations": citations,
        "reference_entries": entries,
        "missing_reference_entries": sorted(set(citations) - set(numbers)),
        "uncited_reference_entries": sorted(set(numbers) - set(citations)),
        "duplicate_reference_numbers": sorted(number for number, count in counts.items() if count > 1),
    }


def renumber(document: Document, rules: dict) -> int:
    entries = reference_entries(document, rules)
    mapping = {entry["number"]: index for index, entry in enumerate(entries, start=1)}
    if len(mapping) != len(entries):
        raise ValueError("Cannot renumber duplicate reference labels; resolve duplicates first.")
    for entry, new_number in zip(entries, range(1, len(entries) + 1)):
        paragraph = document.paragraphs[entry["paragraph"] - 1]
        changed = False
        for run in paragraph.runs:
            match = ENTRY.match(run.text)
            if match:
                run.text = f"[{new_number}] {match.group(2)}"
                changed = True
                break
        if not changed:
            raise ValueError(f"Reference entry {entry['number']} spans multiple Word runs and cannot be renumbered safely.")

    def replacement(match: re.Match) -> str:
        tokens = re.findall(r"\d+|[,，、\-–\s]+", match.group(2))
        return match.group(1) + "".join(str(mapping.get(int(token), int(token))) if token.isdigit() else token for token in tokens) + match.group(3)

    for paragraph in all_paragraphs(document):
        for run in paragraph.runs:
            if BRACKETED.search(run.text):
                run.text = BRACKETED.sub(replacement, run.text)
    return len(entries)


def insert_authoritative_entries(document: Document, rules: dict, metadata: dict[str, str]) -> int:
    report = audit(document, rules)
    missing = report["missing_reference_entries"]
    entries = reference_entries(document, rules)
    if not missing:
        return 0
    unavailable = [number for number in missing if str(number) not in metadata]
    if unavailable:
        raise ValueError("Missing authoritative metadata for citation number(s): " + ", ".join(map(str, unavailable)))
    heading_index = next((index for index, paragraph in enumerate(document.paragraphs) if paragraph_role(paragraph, rules) == "reference_heading"), None)
    if heading_index is None:
        raise ValueError("Cannot insert references because no 参考文献/References heading was found.")
    anchor = document.paragraphs[entries[-1]["paragraph"] - 1] if entries else document.paragraphs[heading_index]
    style = anchor.style
    for number in missing:
        paragraph = document.add_paragraph(style=style)
        paragraph.add_run(f"[{number}] {metadata[str(number)]}")
        anchor._p.addnext(paragraph._p)
        anchor = paragraph
    return len(missing)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("--rules", required=True, type=Path)
    parser.add_argument("--report", type=Path, help="Write the citation/reference audit JSON.")
    parser.add_argument("--renumber", action="store_true", help="Renumber existing numeric entries and citations; does not create sources.")
    parser.add_argument("--metadata-json", type=Path, help="Authoritative JSON object: {citation_number: formatted_reference_entry}.")
    parser.add_argument("--output", type=Path, help="Required with --renumber or --metadata-json.")
    args = parser.parse_args()
    document, rules = Document(args.input_docx), read_rules(args.rules)
    report = audit(document, rules)
    payload = json.dumps(report, ensure_ascii=False, indent=2)
    if args.report:
        args.report.write_text(payload + "\n", encoding="utf-8")
        print(f"Wrote reference audit: {args.report}")
    else:
        print(payload)
    if args.renumber or args.metadata_json:
        if not args.output:
            parser.error("--output is required with --renumber or --metadata-json")
    if args.metadata_json:
        metadata = json.loads(args.metadata_json.read_text(encoding="utf-8"))
        if not isinstance(metadata, dict) or not all(isinstance(key, str) and isinstance(value, str) for key, value in metadata.items()):
            parser.error("--metadata-json must be an object mapping citation numbers to formatted reference text.")
        try:
            inserted = insert_authoritative_entries(document, rules, metadata)
        except ValueError as error:
            raise SystemExit(str(error)) from error
        print(f"Inserted {inserted} authoritative reference entry/entries.")
        report = audit(document, rules)
    if args.renumber:
        if report["missing_reference_entries"] or report["duplicate_reference_numbers"]:
            raise SystemExit("Refusing to renumber: resolve missing or duplicate reference labels first.")
        print(f"Renumbered {renumber(document, rules)} existing entry/entries.")
    if args.renumber or args.metadata_json:
        document.save(args.output)
        print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
