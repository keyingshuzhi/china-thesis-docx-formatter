"""Shared DOCX extraction and conservative formatting helpers."""
from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path
from typing import Any

from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RULES_SCHEMA_VERSION = "1.0"

# These profiles deliberately describe *what to check*, not how a university
# wants it typeset.  The template remains the only source of font, margin, and
# page-break rules.
DEGREE_PROFILES: dict[str, dict[str, Any]] = {
    "auto": {
        "label": "中国学位论文（自动识别）",
        "review_sections": ["abstract_cn", "abstract_en", "keywords_cn", "keywords_en", "toc", "references"],
    },
    "undergraduate": {
        "label": "中国本科毕业论文",
        "review_sections": ["abstract_cn", "abstract_en", "keywords_cn", "keywords_en", "toc", "references", "acknowledgements", "appendices"],
    },
    "master": {
        "label": "中国硕士学位论文",
        "review_sections": ["declaration", "abstract_cn", "abstract_en", "keywords_cn", "keywords_en", "toc", "references", "acknowledgements", "appendices", "research_achievements"],
    },
    "doctoral": {
        "label": "中国博士学位论文",
        "review_sections": ["declaration", "abstract_cn", "abstract_en", "keywords_cn", "keywords_en", "toc", "references", "acknowledgements", "appendices", "research_achievements"],
    },
}

SECTION_LABELS = {
    "cover": "封面/题名页", "declaration": "原创性声明或授权页", "abstract_cn": "中文摘要",
    "abstract_en": "英文摘要", "keywords_cn": "中文关键词", "keywords_en": "英文关键词",
    "toc": "目录", "references": "参考文献", "acknowledgements": "致谢",
    "appendices": "附录", "research_achievements": "攻读学位期间成果页",
}


def _optional_pt(value: Any) -> float | None:
    return None if value is None else round(float(value.pt), 2)


def _optional_cm(value: Any) -> float | None:
    return None if value is None else round(float(value.cm), 3)


def _enum_name(value: Any) -> str | None:
    return getattr(value, "name", None) if value is not None else None


def _east_asia_font(element: Any) -> str | None:
    r_pr = getattr(element, "rPr", None)
    if r_pr is None:
        return None
    fonts = r_pr.find(qn("w:rFonts"))
    return fonts.get(qn("w:eastAsia")) if fonts is not None else None


def _color_value(font: Any) -> str | None:
    rgb = getattr(getattr(font, "color", None), "rgb", None)
    return str(rgb) if rgb is not None else None


def font_rule(source: Any) -> dict[str, Any]:
    # Styles and runs expose ``font`` directly; paragraphs carry it on runs.
    run = source if hasattr(source, "font") else next((item for item in source.runs if item.text), None)
    if run is None:
        return {"name": None, "east_asia": None, "size_pt": None, "bold": None,
                "italic": None, "underline": None, "color": None}
    font = run.font
    return {
        "name": font.name,
        "east_asia": _east_asia_font(run._element if hasattr(run, "_element") else run.element),
        "size_pt": _optional_pt(font.size),
        "bold": font.bold,
        "italic": font.italic,
        "underline": bool(font.underline) if font.underline is not None else None,
        "color": _color_value(font),
    }


def paragraph_rule(source: Any) -> dict[str, Any]:
    fmt = source.paragraph_format
    line = fmt.line_spacing
    return {
        "alignment": _enum_name(fmt.alignment),
        "first_line_indent_cm": _optional_cm(fmt.first_line_indent),
        "left_indent_cm": _optional_cm(fmt.left_indent),
        "right_indent_cm": _optional_cm(fmt.right_indent),
        "space_before_pt": _optional_pt(fmt.space_before),
        "space_after_pt": _optional_pt(fmt.space_after),
        "line_spacing_pt": _optional_pt(line) if hasattr(line, "pt") else None,
        "line_spacing_multiple": round(float(line), 3) if isinstance(line, float) else None,
        "keep_with_next": fmt.keep_with_next,
        "page_break_before": fmt.page_break_before,
    }


def format_rule(source: Any) -> dict[str, Any]:
    return {"font": font_rule(source), "paragraph": paragraph_rule(source)}


def _style_role(style_name: str, style_id: str) -> str | None:
    value = f"{style_name} {style_id}".lower().replace(" ", "")
    if any(token in value for token in ("参考文献", "bibliography", "references")):
        return "reference_entry"
    if any(token in value for token in ("图题", "表题", "caption", "题注")):
        return "caption"
    if "封面" in value:
        return "cover_title"
    if any(token in value for token in ("论文题目", "文章标题", "title")) and "subtitle" not in value:
        return "title"
    if "摘要" in value or "abstract" in value:
        return "abstract"
    for level in range(1, 5):
        if re.search(fr"(?:heading|标题){level}", value) or f"{level}级标题" in value:
            return f"heading_{level}"
    return None


def semantic_section_role(text: str) -> str | None:
    """Identify common Chinese degree-thesis section labels from visible text.

    This is only a structural hint.  It does not infer a mandatory institutional
    component, its order, or its formatting.
    """
    compact = re.sub(r"\s+", "", text).strip()
    lowered = compact.casefold()
    if re.fullmatch(r"(?:原创性声明|独创性声明|学位论文原创性声明|学位论文使用授权书|版权使用授权书)", compact):
        return "declaration"
    if re.fullmatch(r"(?:摘要|中文摘要|内容摘要)", compact):
        return "abstract_cn"
    if lowered in {"abstract", "abstracts"}:
        return "abstract_en"
    if compact.startswith(("关键词", "关键字")):
        return "keywords_cn"
    if lowered.startswith(("keywords", "key words")):
        return "keywords_en"
    if re.fullmatch(r"(?:目录|目次|contents|tableofcontents)", compact, flags=re.IGNORECASE):
        return "toc"
    if re.fullmatch(r"(?:参考文献|references|bibliography)", compact, flags=re.IGNORECASE):
        return "references"
    if re.fullmatch(r"(?:致谢|致謝|acknowledg(?:e)?ments?)", compact, flags=re.IGNORECASE):
        return "acknowledgements"
    if re.fullmatch(r"(?:附录|附錄|appendi(?:x|ces))(?:[A-Z一二三四五六七八九十0-9].*)?", compact, flags=re.IGNORECASE):
        return "appendices"
    if re.search(r"(?:攻读|攻讀).{0,12}(?:学位|學位).{0,16}(?:成果|发表|發表|论文|論文)|(?:发表|發表).{0,12}(?:论文|論文)|学术成果", compact):
        return "research_achievements"
    return None


def _structure_summary(document: Any, degree: str) -> dict[str, Any]:
    detected = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        role = semantic_section_role(paragraph.text)
        if role:
            detected.append({"role": role, "label": SECTION_LABELS[role], "paragraph": index,
                             "text": paragraph.text.strip(), "style": paragraph.style.name if paragraph.style else None})
    found = {item["role"] for item in detected}
    profile = DEGREE_PROFILES[degree]
    return {
        "detected_sections": detected,
        "review_sections": [{"role": role, "label": SECTION_LABELS[role], "detected": role in found}
                            for role in profile["review_sections"]],
        "note": "These are common domestic degree-thesis components for manual confirmation. The official school template and written rules decide applicability, order, and formatting.",
    }


def _style_data(style: Any) -> dict[str, Any]:
    return {
        "style_id": style.style_id,
        "type": _enum_name(style.type),
        "based_on": style.base_style.name if style.base_style is not None else None,
        "format": format_rule(style),
    }


def _page_rule(section: Any) -> dict[str, Any]:
    return {
        "orientation": "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait",
        "page_width_cm": _optional_cm(section.page_width), "page_height_cm": _optional_cm(section.page_height),
        "margin_top_cm": _optional_cm(section.top_margin), "margin_bottom_cm": _optional_cm(section.bottom_margin),
        "margin_left_cm": _optional_cm(section.left_margin), "margin_right_cm": _optional_cm(section.right_margin),
        "header_distance_cm": _optional_cm(section.header_distance), "footer_distance_cm": _optional_cm(section.footer_distance),
    }


def _header_footer_summary(section: Any) -> dict[str, Any]:
    def summary(part: Any) -> dict[str, Any]:
        return {"linked_to_previous": part.is_linked_to_previous, "paragraph_count": len(part.paragraphs),
                "has_visible_text": any(p.text.strip() for p in part.paragraphs)}
    return {"header": summary(section.header), "footer": summary(section.footer)}


def _in_use_style_names(document: Any) -> set[str]:
    names = {p.style.name for p in document.paragraphs if p.style is not None}
    for table in document.tables:
        names.update(p.style.name for row in table.rows for cell in row.cells for p in cell.paragraphs)
    return names


def _table_rule(document: Any) -> dict[str, Any]:
    if not document.tables:
        return {"observed": False}
    styles = Counter(t.style.name if t.style is not None else None for t in document.tables)
    cell = next((c for t in document.tables for row in t.rows for c in row.cells), None)
    return {
        "observed": True, "style_name": styles.most_common(1)[0][0],
        "alignment": _enum_name(document.tables[0].alignment),
        "cell_vertical_alignment": _enum_name(cell.vertical_alignment) if cell else None,
        "cell_text_format": format_rule(cell.paragraphs[0]) if cell and cell.paragraphs else None,
    }


def extract_rules(document: Any, source_name: str, degree: str = "auto") -> dict[str, Any]:
    if degree not in DEGREE_PROFILES:
        raise ValueError(f"Unsupported degree level: {degree}")
    warnings: list[str] = []
    used = _in_use_style_names(document)
    selected = [s for s in document.styles if s.name in used]
    if not selected:
        selected = [s for s in document.styles if s.type == WD_STYLE_TYPE.PARAGRAPH and s.style_id == "Normal"]
    styles = {s.name: _style_data(s) for s in selected}
    role_rules: dict[str, Any] = {}
    for style in selected:
        role = _style_role(style.name, style.style_id)
        if role and role not in role_rules:
            role_rules[role] = {"source_style": style.name, "source_style_id": style.style_id,
                                "confidence": "inferred_from_style_name", "format": format_rule(style)}
    normal = next((s for s in document.styles if s.style_id == "Normal"), None)
    if "body" not in role_rules and normal is not None:
        role_rules["body"] = {"source_style": normal.name, "source_style_id": normal.style_id,
                              "confidence": "normal_style_fallback", "format": format_rule(normal)}
        warnings.append("Body formatting was inferred from Normal; confirm it matches thesis body text.")
    if not any(key.startswith("heading_") for key in role_rules):
        warnings.append("No heading styles were identified; map heading roles before formatting headings.")
    if "reference_entry" not in role_rules:
        warnings.append("No reference-entry style was identified; review bibliography formatting before use.")
    pages = [_page_rule(section) for section in document.sections]
    if len({json.dumps(page, sort_keys=True) for page in pages}) > 1:
        warnings.append("Template has multiple page setups; whole-document formatting uses the first setup.")
    return {
        "schema_version": RULES_SCHEMA_VERSION,
        "source": {"template_filename": source_name, "extractor": "academic-paper-formatter"},
        "page": pages[0] if pages else {},
        "sections": [{"index": i + 1, "page": page, "header_footer": _header_footer_summary(section)}
                     for i, (section, page) in enumerate(zip(document.sections, pages))],
        "styles": styles, "role_rules": role_rules, "table_rules": _table_rule(document),
        "figure_rules": {"caption_role": "caption" if "caption" in role_rules else None},
        "reference_rules": {"entry_role": "reference_entry" if "reference_entry" in role_rules else None},
        "degree_profile": {"level": degree, "label": DEGREE_PROFILES[degree]["label"],
                           "scope": "China domestic degree thesis", "is_formatting_authority": False},
        "document_structure": _structure_summary(document, degree),
        "extraction_warnings": warnings,
    }


def read_rules(path: str | Path) -> dict[str, Any]:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def write_rules(rules: dict[str, Any], path: str | Path) -> None:
    Path(path).write_text(json.dumps(rules, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def _set_east_asia_font(run: Any, name: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        from docx.oxml import OxmlElement
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), name)


def apply_format(paragraph: Any, rule: dict[str, Any] | None) -> None:
    if not rule:
        return
    font = rule.get("font", {})
    for run in paragraph.runs:
        if font.get("name"): run.font.name = font["name"]
        if font.get("east_asia"): _set_east_asia_font(run, font["east_asia"])
        if font.get("size_pt") is not None: run.font.size = Pt(font["size_pt"])
        for key in ("bold", "italic", "underline"):
            if font.get(key) is not None: setattr(run.font, key, font[key])
        if font.get("color"): run.font.color.rgb = RGBColor.from_string(font["color"])
    p_rule = rule.get("paragraph", {})
    alignment = p_rule.get("alignment")
    if alignment and hasattr(WD_ALIGN_PARAGRAPH, alignment): paragraph.alignment = getattr(WD_ALIGN_PARAGRAPH, alignment)
    fmt = paragraph.paragraph_format
    for key, target in (("first_line_indent_cm", "first_line_indent"), ("left_indent_cm", "left_indent"),
                        ("right_indent_cm", "right_indent")):
        if p_rule.get(key) is not None: setattr(fmt, target, Cm(p_rule[key]))
    for key, target in (("space_before_pt", "space_before"), ("space_after_pt", "space_after")):
        if p_rule.get(key) is not None: setattr(fmt, target, Pt(p_rule[key]))
    if p_rule.get("line_spacing_pt") is not None: fmt.line_spacing = Pt(p_rule["line_spacing_pt"])
    elif p_rule.get("line_spacing_multiple") is not None: fmt.line_spacing = p_rule["line_spacing_multiple"]
    for key in ("keep_with_next", "page_break_before"):
        if p_rule.get(key) is not None: setattr(fmt, key, p_rule[key])


def apply_page_rule(document: Any, page: dict[str, Any]) -> None:
    for section in document.sections:
        if page.get("orientation") == "landscape": section.orientation = WD_ORIENT.LANDSCAPE
        elif page.get("orientation") == "portrait": section.orientation = WD_ORIENT.PORTRAIT
        for key, target in (("page_width_cm", "page_width"), ("page_height_cm", "page_height"),
                            ("margin_top_cm", "top_margin"), ("margin_bottom_cm", "bottom_margin"),
                            ("margin_left_cm", "left_margin"), ("margin_right_cm", "right_margin"),
                            ("header_distance_cm", "header_distance"), ("footer_distance_cm", "footer_distance")):
            if page.get(key) is not None: setattr(section, target, Cm(page[key]))


def paragraph_role(paragraph: Any, rules: dict[str, Any]) -> str | None:
    text = paragraph.text.strip()
    if not text: return None
    semantic_role = semantic_section_role(text)
    if semantic_role == "references":
        return "reference_heading"
    if semantic_role:
        return semantic_role
    style_name = paragraph.style.name if paragraph.style is not None else ""
    for role, rule in rules.get("role_rules", {}).items():
        if rule.get("source_style") == style_name: return role
    if re.match(r"^(?:图|表|figure|table)\s*[0-9一二三四五六七八九十]", text, flags=re.IGNORECASE): return "caption"
    if re.match(r"^(?:第[一二三四五六七八九十\d]+[章节]|\d+(?:\.\d+){0,3}\s+)", text):
        return f"heading_{min(text.count('.') + 1, 4)}"
    return "body"


def role_format(rules: dict[str, Any], role: str) -> dict[str, Any] | None:
    return rules.get("role_rules", {}).get(role, {}).get("format")
