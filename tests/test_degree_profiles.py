#!/usr/bin/env python3
"""Regression checks for domestic degree-thesis structure profiling."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))
from docx_rules import extract_rules, paragraph_role, semantic_section_role  # noqa: E402


def main() -> None:
    document = Document()
    for text in ("原创性声明", "摘要", "关键词：论文排版", "Abstract", "Keywords: thesis formatting", "目录", "参考文献", "致谢", "附录 A", "攻读博士学位期间发表的论文"):
        document.add_paragraph(text)
    rules = extract_rules(document, "doctoral-template.docx", "doctoral")
    assert rules["degree_profile"]["level"] == "doctoral"
    detected = {item["role"] for item in rules["document_structure"]["detected_sections"]}
    expected = {"declaration", "abstract_cn", "keywords_cn", "abstract_en", "keywords_en", "toc", "references", "acknowledgements", "appendices", "research_achievements"}
    assert expected <= detected
    assert semantic_section_role("参考文献") == "references"
    assert paragraph_role(document.paragraphs[6], rules) == "reference_heading"
    print("degree profile checks passed")


if __name__ == "__main__":
    main()
