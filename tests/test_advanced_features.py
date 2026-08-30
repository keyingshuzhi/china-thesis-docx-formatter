"""Smoke tests for safe completion, reference, field, PDF, and anchor helpers."""
from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree
from reportlab.pdfgen import canvas
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))
from normalize_floating_figures import NS, configure_anchor  # noqa: E402


def make_document(path: Path) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "STSong")
    document.add_paragraph("正文 [[TODO: 摘要]]")
    document.add_paragraph("引用见[1]")
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 已有来源")
    document.save(path)


def main() -> None:
    python = sys.executable
    with tempfile.TemporaryDirectory() as temp_dir:
        work = Path(temp_dir)
        source = work / "thesis.docx"
        make_document(source)
        rules = {"schema_version": "1.0", "source": {"template_filename": "test.docx"}, "page": {}, "styles": {},
                 "role_rules": {"body": {"format": {} }}, "table_rules": {"observed": False}, "extraction_warnings": []}
        rule_path = work / "rules.json"
        rule_path.write_text(json.dumps(rules), encoding="utf-8")
        placeholders, replacements, completed = work / "placeholders.json", work / "replacements.json", work / "completed.docx"
        subprocess.run([python, str(SCRIPTS / "complete_content.py"), str(source), "--report", str(placeholders)], check=True)
        assert json.loads(placeholders.read_text(encoding="utf-8"))["placeholders"]
        replacements.write_text(json.dumps({"[[TODO: 摘要]]": "已审阅摘要"}, ensure_ascii=False), encoding="utf-8")
        subprocess.run([python, str(SCRIPTS / "complete_content.py"), str(source), "--replacements", str(replacements), "--output", str(completed)], check=True)
        assert "已审阅摘要" in "\n".join(paragraph.text for paragraph in Document(completed).paragraphs)
        reference_report = work / "references.json"
        subprocess.run([python, str(SCRIPTS / "manage_references.py"), str(source), "--rules", str(rule_path), "--report", str(reference_report)], check=True)
        assert json.loads(reference_report.read_text(encoding="utf-8"))["missing_reference_entries"] == []
        renumbered = work / "renumbered.docx"
        subprocess.run([python, str(SCRIPTS / "manage_references.py"), str(source), "--rules", str(rule_path), "--renumber", "--output", str(renumbered)], check=True)
        assert "[1] 已有来源" in "\n".join(paragraph.text for paragraph in Document(renumbered).paragraphs)
        missing_source = work / "missing-reference.docx"
        missing_document = Document(source)
        missing_document.add_paragraph("另见[2]")
        missing_document.save(missing_source)
        metadata = work / "metadata.json"
        metadata.write_text(json.dumps({"2": "权威来源条目"}, ensure_ascii=False), encoding="utf-8")
        repaired = work / "repaired.docx"
        subprocess.run([python, str(SCRIPTS / "manage_references.py"), str(missing_source), "--rules", str(rule_path), "--metadata-json", str(metadata), "--output", str(repaired)], check=True)
        assert "[2] 权威来源条目" in "\n".join(paragraph.text for paragraph in Document(repaired).paragraphs)
        fields = work / "fields.docx"
        subprocess.run([python, str(SCRIPTS / "update_fields.py"), str(completed), str(fields)], check=True)
        with zipfile.ZipFile(fields) as archive:
            assert b"updateFields" in archive.read("word/settings.xml")
        pdf = work / "guide.pdf"
        report = canvas.Canvas(str(pdf))
        report.drawString(72, 720, "Font 12 pt. Margin 2.5 cm. Reference list follows the guide.")
        report.save()
        pdf_rules = work / "pdf-rules.json"
        subprocess.run([python, str(SCRIPTS / "parse_pdf_rules.py"), str(pdf), "--output", str(pdf_rules)], check=True)
        assert json.loads(pdf_rules.read_text(encoding="utf-8"))["requirements"]
    anchor = etree.Element("{%s}anchor" % NS["wp"])
    etree.SubElement(anchor, "{%s}positionH" % NS["wp"])
    etree.SubElement(anchor, "{%s}positionV" % NS["wp"])
    configure_anchor(anchor, "center", "top", "paragraph")
    assert anchor.get("allowOverlap") == "0"
    assert anchor.find("wp:positionH/wp:align", NS).text == "center"


if __name__ == "__main__":
    main()
