#!/usr/bin/env python3
"""Create a fictional, privacy-safe fixture set for the formatter skill."""
from __future__ import annotations

import argparse
import json
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw
from lxml import etree
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def set_east_asia(style, font_name: str) -> None:
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)


def add_field(paragraph, instruction: str, display: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend((begin, instruction_text, separate, text, end))


def add_bookmark(paragraph, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "1")
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "1")
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def configure_styles(document: Document, template: bool) -> None:
    normal = document.styles["Normal"]
    normal.font.name, normal.font.size = "Hiragino Sans GB", Pt(12 if template else 10.5)
    set_east_asia(normal, "Hiragino Sans GB")
    normal.paragraph_format.first_line_indent = Cm(0.74 if template else 0)
    normal.paragraph_format.line_spacing = 1.5
    title = document.styles.add_style("论文题目", WD_STYLE_TYPE.PARAGRAPH)
    title.font.name, title.font.size, title.font.bold = "Hiragino Sans GB", Pt(18), True
    set_east_asia(title, "Hiragino Sans GB")
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading = document.styles["Heading 1"]
    heading.font.name, heading.font.size, heading.font.bold = "Hiragino Sans GB", Pt(16), True
    set_east_asia(heading, "Hiragino Sans GB")
    caption = document.styles["Caption"]
    caption.font.name, caption.font.size = "Hiragino Sans GB", Pt(10.5)
    set_east_asia(caption, "Hiragino Sans GB")
    reference = document.styles.add_style("参考文献", WD_STYLE_TYPE.PARAGRAPH)
    reference.font.name, reference.font.size = "Hiragino Sans GB", Pt(10.5)
    set_east_asia(reference, "Hiragino Sans GB")


def set_page(document: Document, template: bool) -> None:
    section = document.sections[0]
    if template:
        section.top_margin, section.bottom_margin = Cm(2.5), Cm(2.0)
        section.left_margin, section.right_margin = Cm(3.0), Cm(2.5)
    else:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(2.0)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Sample University Thesis - Test Data")


def make_template(path: Path) -> None:
    document = Document()
    configure_styles(document, template=True)
    set_page(document, template=True)
    document.add_paragraph("Sample University Thesis Format", style="论文题目")
    document.add_paragraph("1 Introduction", style="Heading 1")
    document.add_paragraph("This paragraph demonstrates the body style.", style="Normal")
    document.add_paragraph("Figure 1 System flow", style="Caption")
    document.add_paragraph("References")
    document.add_paragraph("[1] Sample Author. Sample Reference[J]. Sample Journal, 2025.", style="参考文献")
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text, table.cell(0, 1).text = "Field", "Description"
    table.cell(1, 0).text, table.cell(1, 1).text = "Sample", "Automation fixture"
    document.save(path)


def make_thesis(path: Path, figure: Path) -> None:
    document = Document()
    configure_styles(document, template=False)
    set_page(document, template=False)
    document.add_paragraph("Template-driven Thesis Formatting Verification", style="论文题目")
    document.add_paragraph("Table of Contents")
    add_field(document.add_paragraph(), ' TOC \\o "1-3" \\h \\z \\u ', "右键更新目录")
    document.add_paragraph("Abstract", style="Heading 1")
    document.add_paragraph("[[TODO:abstract]]", style="Normal")
    document.add_paragraph("1 Introduction", style="Heading 1")
    document.add_paragraph("This fixture validates template extraction, paragraph formatting, and reference auditing[1, 2].", style="Normal")
    figure_caption = document.add_paragraph("Figure 1 Test flow", style="Caption")
    add_bookmark(figure_caption, "_RefTest")
    document.add_picture(str(figure), width=Cm(7.0))
    document.add_paragraph("The flow covers rule extraction, formatting, and quality auditing.", style="Normal")
    document.add_paragraph("Table 1 Test coverage", style="Caption")
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text, table.cell(0, 1).text = "Feature", "Expected result"
    table.cell(1, 0).text, table.cell(1, 1).text = "Completion", "Approved text is inserted"
    table.cell(2, 0).text, table.cell(2, 1).text = "Reference audit", "Missing entry [2] is found"
    document.add_paragraph("References")
    document.add_paragraph("[1] Wang Ming. Template-driven Document Processing[J]. Document Engineering, 2025, 10(1): 1-8.", style="参考文献")
    crossref = document.add_paragraph("Cross-reference example: ")
    add_field(crossref, " REF _RefTest \\h ", "Figure 1")
    document.save(path)
    make_first_image_floating(path)


def make_first_image_floating(path: Path) -> None:
    """Convert one valid inline picture to a minimal, valid floating anchor."""
    with tempfile.TemporaryDirectory() as temp_dir:
        work = Path(temp_dir)
        with zipfile.ZipFile(path) as archive:
            archive.extractall(work)
        xml_path = work / "word" / "document.xml"
        tree = etree.parse(str(xml_path))
        inline = tree.find(f".//{{{WP}}}inline")
        if inline is not None:
            inline.tag = f"{{{WP}}}anchor"
            inline.attrib.clear()
            inline.attrib.update({"distT": "0", "distB": "0", "distL": "114300", "distR": "114300", "simplePos": "0",
                                 "relativeHeight": "251658240", "behindDoc": "0", "locked": "0", "layoutInCell": "1", "allowOverlap": "1"})
            simple = etree.Element(f"{{{WP}}}simplePos", x="0", y="0")
            position_h = etree.Element(f"{{{WP}}}positionH", relativeFrom="column")
            etree.SubElement(position_h, f"{{{WP}}}align").text = "left"
            position_v = etree.Element(f"{{{WP}}}positionV", relativeFrom="paragraph")
            etree.SubElement(position_v, f"{{{WP}}}posOffset").text = "0"
            wrap = etree.Element(f"{{{WP}}}wrapSquare", wrapText="bothSides")
            inline.insert(0, simple)
            inline.insert(1, position_h)
            inline.insert(2, position_v)
            inline.insert(4, wrap)
            tree.write(str(xml_path), encoding="UTF-8", xml_declaration=True, standalone=True)
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
            for file_path in work.rglob("*"):
                if file_path.is_file():
                    archive.write(file_path, file_path.relative_to(work).as_posix())


def make_pdf(path: Path) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    report = canvas.Canvas(str(path))
    report.setFont("STSong-Light", 14)
    report.drawString(72, 760, "示例大学毕业论文排版规范（虚构测试文件）")
    report.setFont("STSong-Light", 11)
    lines = [
        "1. 正文使用宋体小四，首行缩进 2 个字符，1.5 倍行距。",
        "2. 页边距：上 2.5 cm，下 2.0 cm，左 3.0 cm，右 2.5 cm。",
        "3. 一级标题使用黑体 16 pt，图表题注使用宋体 10.5 pt。",
        "4. 目录、页码和交叉引用在提交前必须更新。",
        "5. 参考文献采用数字顺序编码，文内引文必须与条目对应。",
    ]
    for index, line in enumerate(lines):
        report.drawString(72, 710 - index * 32, line)
    report.showPage()
    report.setFont("STSong-Light", 11)
    report.drawString(72, 760, "附加说明：图像不得遮挡正文；表格内容应完整可读。")
    report.save()


def make_figure(path: Path) -> None:
    image = Image.new("RGB", (800, 360), "white")
    drawing = ImageDraw.Draw(image)
    drawing.rounded_rectangle((50, 110, 230, 240), radius=16, outline="#28527a", width=5, fill="#eaf3f8")
    drawing.rounded_rectangle((310, 110, 490, 240), radius=16, outline="#28527a", width=5, fill="#eaf3f8")
    drawing.rounded_rectangle((570, 110, 750, 240), radius=16, outline="#28527a", width=5, fill="#eaf3f8")
    drawing.line((230, 175, 310, 175), fill="#28527a", width=5)
    drawing.line((490, 175, 570, 175), fill="#28527a", width=5)
    image.save(path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", type=Path, default=Path("test_data"))
    args = parser.parse_args()
    output = args.output_dir
    output.mkdir(parents=True, exist_ok=True)
    figure = output / "figure-placeholder.png"
    make_figure(figure)
    make_template(output / "school-template.docx")
    make_thesis(output / "thesis-source.docx", figure)
    make_pdf(output / "school-guideline.pdf")
    (output / "approved-content.json").write_text(json.dumps({"[[TODO:abstract]]": "This privacy-safe fixture validates an end-to-end thesis-formatting workflow."}, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    (output / "verified-references.json").write_text(json.dumps({"2": "Li Hua. Automated Validation of Academic Formatting Rules[M]. Example Press, 2024."}, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    (output / "README.md").write_text("""# 学术论文格式化测试包

所有学校、作者、文献与内容均为虚构测试数据。

1. `school-template.docx`：学校格式模板。
2. `thesis-source.docx`：待处理论文；包含错误页边距、待补全摘要、[2] 缺失文献、TOC/REF 字段和插图。
3. `school-guideline.pdf`：PDF 版格式规范。
4. `approved-content.json`：已审核的占位符替换文本。
5. `verified-references.json`：可补入缺失 [2] 条目的权威元数据。

请按项目根目录的 `SKILL.md` 顺序执行验证。所有输出请写到此目录下的 `output/` 子目录，勿覆盖源数据。
""", encoding="utf-8")
    print(f"Created test data in {output}")


if __name__ == "__main__":
    main()
