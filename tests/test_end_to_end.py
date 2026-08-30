"""Smoke test: a template becomes JSON rules that can format a thesis copy."""
from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"


def east_asia(style, name: str) -> None:
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)


def make_template(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin, section.bottom_margin = Cm(2.5), Cm(2.0)
    section.left_margin, section.right_margin = Cm(3.0), Cm(2.5)
    normal = document.styles["Normal"]
    normal.font.name, normal.font.size = "Times New Roman", Pt(12)
    # Use fonts present in the CI renderer; the extractor itself preserves any
    # font name found in a real school template, including 宋体 and 黑体.
    east_asia(normal, "STSong")
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.5
    heading = document.styles["Heading 1"]
    heading.font.name, heading.font.size, heading.font.bold = "Arial", Pt(16), True
    east_asia(heading, "Heiti SC")
    caption = document.styles["Caption"]
    caption.font.name, caption.font.size = "Times New Roman", Pt(10.5)
    east_asia(caption, "STSong")
    references = document.styles.add_style("参考文献", WD_STYLE_TYPE.PARAGRAPH)
    references.font.name, references.font.size = "Times New Roman", Pt(10.5)
    east_asia(references, "STSong")
    document.add_paragraph("正文样例", style="Normal")
    document.add_paragraph("1 绪论", style="Heading 1")
    document.add_paragraph("图 1 研究流程", style="Caption")
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 示例文献", style="参考文献")
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text, table.cell(0, 1).text = "项目", "数值"
    document.save(path)


def main() -> None:
    python = sys.executable
    with tempfile.TemporaryDirectory() as temp_dir:
        work = Path(temp_dir)
        template, source, rules, output = (work / name for name in ("school-template.docx", "thesis.docx", "rules.json", "formatted.docx"))
        make_template(template)
        make_template(source)
        subprocess.run([python, str(SCRIPTS / "analyze_docx.py"), str(template), "--output", str(rules)], check=True)
        parsed = json.loads(rules.read_text(encoding="utf-8"))
        assert parsed["source"]["template_filename"] == "school-template.docx"
        assert parsed["role_rules"]["body"]["format"]["font"]["east_asia"] == "STSong"
        subprocess.run([python, str(SCRIPTS / "format_document.py"), str(source), str(output), "--rules", str(rules)], check=True)
        subprocess.run([python, str(SCRIPTS / "validate_format.py"), str(rules), "--document", str(output)], check=True)


if __name__ == "__main__":
    main()
